"""
MockMate Profile Service — Port 8002
Resume handling (text extraction), GitHub analysis, skills extraction.
"""
from fastapi import FastAPI, UploadFile, File, HTTPException, Request, Depends
from fastapi.responses import JSONResponse
from pydantic import BaseModel
from typing import Optional, List
import os, io, uuid, httpx, traceback
from pathlib import Path
from dotenv import load_dotenv

load_dotenv(Path(__file__).parent.parent.parent / ".env")

app = FastAPI(title="MockMate Profile Service", version="2.0.0")

GITHUB_TOKEN     = os.getenv("GITHUB_TOKEN", "")
AUTH_SERVICE_URL = os.getenv("AUTH_SERVICE_URL", "http://localhost:8001")

from common.auth import get_current_user_id


@app.exception_handler(Exception)
async def global_exception_handler(request: Request, exc: Exception):
    tb = traceback.format_exc()
    return JSONResponse(status_code=500, content={"error": str(exc), "type": type(exc).__name__, "traceback": tb})


# ─── Health ───────────────────────────────────────────────────────────────────
@app.get("/health")
def health():
    return {"status": "ok", "service": "profile"}


# ─── Resume Upload + Text Extraction ─────────────────────────────────────────
def _extract_text_from_pdf(content: bytes) -> str:
    """Extract text from PDF bytes. Tries pdfplumber first, pypdfium2 as fallback."""
    text = ""

    # Strategy 1: pdfplumber (best for standard text PDFs)
    try:
        import pdfplumber
        with pdfplumber.open(io.BytesIO(content)) as pdf:
            pages_text = []
            for i, page in enumerate(pdf.pages):
                page_text = page.extract_text() or ""
                print(f"[Profile] PDF page {i+1}: {len(page_text)} chars extracted")
                pages_text.append(page_text)
        text = "\n".join(pages_text).strip()
    except Exception as e:
        print(f"[Profile] pdfplumber failed: {e}")

    if len(text) >= 50:
        return text

    # Strategy 2: pypdfium2 fallback (handles some PDFs pdfplumber misses)
    try:
        import pypdfium2 as pdfium
        doc = pdfium.PdfDocument(io.BytesIO(content))
        pages_text = []
        for i in range(len(doc)):
            page = doc[i]
            textpage = page.get_textpage()
            page_text = textpage.get_text_range() or ""
            print(f"[Profile] pypdfium2 page {i+1}: {len(page_text)} chars")
            pages_text.append(page_text)
        fallback_text = "\n".join(pages_text).strip()
        if len(fallback_text) > len(text):
            text = fallback_text
    except Exception as e:
        print(f"[Profile] pypdfium2 failed: {e}")

    print(f"[Profile] Total PDF text extracted: {len(text)} chars")
    return text


def _extract_text_from_docx(content: bytes) -> str:
    """Extract text from DOCX bytes using python-docx."""
    try:
        from docx import Document
        doc = Document(io.BytesIO(content))
        paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]
        return "\n".join(paragraphs).strip()
    except Exception as e:
        print(f"[Profile] DOCX extraction error: {e}")
        return ""


@app.post("/resume/upload")
async def upload_resume(
    request: Request,
    file: UploadFile = File(...),
    user_id: str = Depends(get_current_user_id),
):
    content = await file.read()
    filename = file.filename or "resume"
    ext = filename.rsplit(".", 1)[-1].lower() if "." in filename else "pdf"

    if ext == "pdf":
        resume_text = _extract_text_from_pdf(content)
    elif ext in ("docx", "doc"):
        resume_text = _extract_text_from_docx(content)
    elif ext == "txt":
        resume_text = content.decode("utf-8", errors="ignore")
    else:
        raise HTTPException(400, f"Unsupported file type: .{ext}. Please upload PDF, DOCX, or TXT.")

    if not resume_text or len(resume_text.strip()) < 50:
        char_count = len(resume_text.strip()) if resume_text else 0
        ext_display = ext.upper()
        if ext == "pdf" and char_count == 0:
            detail = (
                "Your PDF appears to be a scanned image or image-only PDF — "
                "no text layer was found. Please export your resume as a text-based PDF "
                "from Word/Google Docs, or upload a TXT file instead."
            )
        elif char_count < 50:
            detail = (
                f"Only {char_count} characters were extracted from your {ext_display} file "
                "(minimum 50 required). Please ensure your resume has readable text content."
            )
        else:
            detail = "Could not extract meaningful text. Please try a PDF, DOCX, or TXT file."
        print(f"[Profile] Upload rejected: {detail}")
        raise HTTPException(422, detail)

    print(f"[Profile] Extracted {len(resume_text)} chars from {filename} for user {user_id}")

    # Forward the original Bearer token so auth service can validate it
    auth_header = request.headers.get("Authorization", "")
    async with httpx.AsyncClient(timeout=30.0) as client:
        resp = await client.put(
            f"{AUTH_SERVICE_URL}/users/me/profile",
            json={"resume_text": resume_text},
            headers={"Authorization": auth_header},
        )
        if resp.status_code not in (200, 201):
            print(f"[Profile] Failed to save resume: {resp.status_code} {resp.text[:200]}")
            raise HTTPException(500, "Failed to save resume to profile.")

    return {
        "status": "ok",
        "filename": filename,
        "characters_extracted": len(resume_text),
        "preview": resume_text[:300] + "..." if len(resume_text) > 300 else resume_text,
    }


@app.post("/resume/extract-text")
async def extract_text_only(file: UploadFile = File(...)):
    """Extract text from a resume file without saving — for preview purposes."""
    content = await file.read()
    filename = file.filename or "resume"
    ext = filename.rsplit(".", 1)[-1].lower() if "." in filename else "pdf"

    if ext == "pdf":
        text = _extract_text_from_pdf(content)
    elif ext in ("docx", "doc"):
        text = _extract_text_from_docx(content)
    elif ext == "txt":
        text = content.decode("utf-8", errors="ignore")
    else:
        raise HTTPException(400, f"Unsupported file type: .{ext}")

    return {"text": text, "characters": len(text)}


# ─── GitHub Analysis ──────────────────────────────────────────────────────────
@app.get("/github/{username}")
async def analyze_github(username: str):
    headers = {"Accept": "application/vnd.github.v3+json"}
    if GITHUB_TOKEN:
        headers["Authorization"] = f"token {GITHUB_TOKEN}"

    async with httpx.AsyncClient(timeout=15.0) as client:
        try:
            resp = await client.get(
                f"https://api.github.com/users/{username}/repos",
                headers=headers, params={"sort": "updated", "per_page": 30},
            )
            if resp.status_code != 200:
                raise HTTPException(404, f"GitHub user '{username}' not found")
            repos = resp.json()
        except httpx.RequestError:
            raise HTTPException(503, "Cannot reach GitHub API")

    lang_counts: dict = {}
    notable: list = []
    for repo in repos:
        lang = repo.get("language")
        if lang:
            lang_counts[lang] = lang_counts.get(lang, 0) + repo.get("stargazers_count", 0) + 1
        if repo.get("stargazers_count", 0) > 2:
            notable.append(repo["name"])

    top_langs = sorted(lang_counts, key=lang_counts.get, reverse=True)[:6]
    return {
        "username": username,
        "top_languages": top_langs,
        "public_repos": len(repos),
        "notable_repos": notable[:5],
        "contribution_score": min(len(repos) * 3, 100),
    }


# ─── Skills Extraction ────────────────────────────────────────────────────────
TECH_VOCAB = [
    "python","java","javascript","typescript","go","rust","c++","c#","kotlin","swift",
    "react","vue","angular","next.js","node.js","fastapi","django","flask","spring",
    "aws","gcp","azure","docker","kubernetes","terraform","linux",
    "postgresql","mysql","mongodb","redis","elasticsearch","kafka","rabbitmq",
    "graphql","rest","grpc","machine learning","deep learning","pytorch","tensorflow",
    "scikit-learn","pandas","numpy","nlp","llm","transformers","hugging face",
    "git","ci/cd","microservices","system design","data structures","algorithms",
    "computer vision","reinforcement learning","mlops","spark","hadoop","databricks",
]

@app.post("/skills/extract")
def extract_skills(payload: dict):
    text  = payload.get("text", "").lower()
    found = [t for t in TECH_VOCAB if t in text]
    return {"skills": found, "count": len(found)}
